'''
Created on Feb 6, 2024

@author: Ngô Anh Tuấn
'''
from _datetime import datetime
import json
import os
import re

from django.contrib.auth.decorators import login_required, permission_required
from django.http.response import HttpResponse
from django.shortcuts import render, redirect, resolve_url
from django.template.context import RequestContext
from django.views.decorators.http import require_http_methods
from docx import Document

from Quiz.models.Answer import Answer
from Quiz.models.ExamQuestion import ExamQuestion
from Quiz.models.Question import Question
from Quiz.models.Subject import Subject
from Quiz.test import BASE_DIR
from Quiz.util.DateEncoder import DateTimeEncoder
from QuizWebapp import settings


@login_required()
def question_import(request):
    context = {}
    try:
        subjects = Subject.objects.all()
        context.update({'subjects':subjects})
        if request.method == 'POST':
            subject_id = request.POST.get('slSubject')
            print(subject_id)
            strTime = datetime.now().strftime('%d%m%Y_%H%M%S')
            handle_uploaded_file(request.FILES["templateFile"],strTime,"template",subject_id)
            images = request.FILES.getlist("listImageFile")
            for image in images:
                handle_uploaded_file(image,strTime,"image",subject_id)
            return redirect(resolve_url('quiz:question'))
    except Exception as ex:
        context.update({'has_error':str(ex)})
    return render(request, "quiz/question/question-import.html", context)
@login_required()
@permission_required('Quiz.import_question', login_url= 'quiz:permission_error')
def index(request):
    context = {}
    subject_id = request.GET.get('subject_id')
    subjects = Subject.objects.all()
    # subject
    if subject_id is None:
        subject_id = subjects.first().id
    questions = Question.objects.filter(subject_id = subject_id)
    context.update({'subjects':subjects,'questions':questions})
    context.update({'STATIC_URL':settings.STATIC_URL})
    return render(request, "quiz/question/question-index.html", context)
@login_required()
@require_http_methods(["POST", ])
def get_question_answer(request):
    try:
        question_id  = request.POST['question_id']
        answers = Answer.objects.filter(question_id = question_id)
        return HttpResponse(json.dumps({'handle':'success', 'answers':list(answers)}, cls=DateTimeEncoder) , content_type="application/json")
    except Exception as ex:
        return HttpResponse(json.dumps({'handle':'error', "msg": str(ex)}), content_type="application/json")
@login_required()
@require_http_methods(["POST", ])
def load_subject_question(request):
    try:
        
        subject_id  = request.POST['subject_id']
        questions = Question.objects.filter(subject_id = subject_id)
        return HttpResponse(json.dumps({'handle':'success', 'questions':list(questions)}, cls=DateTimeEncoder) , content_type="application/json")
    except Exception as ex:
        return HttpResponse(json.dumps({'handle':'error', "msg": str(ex)}), content_type="application/json")
@login_required()
def delete(request,question_id):
    try:
        # exam question
        exam_questions = ExamQuestion.objects.filter(question_id = question_id)
        exam_questions.delete()
        # answer
        answers = Answer.objects.filter(question_id = question_id)
        answers.delete()
        # question
        question = Question.objects.get(id=int(question_id))
        question.delete()
        # check permission error
        return HttpResponse(json.dumps({'handle':'success',}), content_type="application/json");
    except Exception as ex:
        return HttpResponse(json.dumps({'handle':'error', "msg": str(ex)}), content_type="application/json")
def handle_uploaded_file(f,strTime,strType,subject_id):
    folder_path = os.path.join(BASE_DIR, 'Quiz','static','upload',strTime)
    if strType=='image':
        folder_path = os.path.join(folder_path,'image')
    if os.path.isdir(folder_path)== False:
        os.makedirs(folder_path)
    with open( os.path.join(folder_path,f.name), 'wb+') as destination:
        for chunk in f.chunks():
            destination.write(chunk)
    if strType=='template':
        readDocx(os.path.join(folder_path,f.name),subject_id,strTime )
def readDocx(file_name,subject_id,strTime):
    questions = []
    lines = list()
    fname, file_extension = os.path.splitext(file_name)
    if file_extension!= '.docx':
        raise Exception("Sorry, incorrect file format")
    doc = Document(file_name)
    #đọc từng dòng văn bản
    for paragraph in doc.paragraphs:
        #lưu nội dung từng dòng thành 1 list
        data = paragraph.text.split(':')
        #lấy value để insert vào data
        #xóa bỏ khoảng trắng trước dữ liệu
        data_insert = ''.join(dt for dt in data[-1] if not dt.isspace())
       
        lines.append(data_insert)    
    #đọc từng table câu hỏi
    for table in doc.tables:
        question_text = table.rows[0].cells[1].text
        image_name = str(get_image_name(question_text))
        image_name = 'upload/'+strTime+'/image/'+image_name.replace("['", '').replace("']", '').strip()
        print(image_name)
        #lấy ra nội dung câu hỏi
        question = Question()
        question.subject_id = subject_id
        question.image_name = image_name
        question.text = question_text
        print('question: %s'%question_text)
        
        #lấy ra mark
        mark = table.rows[len(table.rows) - 3].cells[1].text
        question.mark  = float(mark)
        #lấy ra unit
        unit = table.rows[len(table.rows) - 2].cells[1].text
        question.unit = unit
        #lấy ra mix choices:
        mix_choices = table.rows[len(table.rows) - 1].cells[1].text
        if mix_choices=='Yes':
            question.is_mix =True
        else:
            question.is_mix = False
        #lấy ra đáp án
        correct_answer = table.rows[len(table.rows) - 4].cells[1].text
        #for i in range (1, len(table.rows) - 4):
            #if table.rows[i].cell[0].text == correct_answer:
                #answer = tab.rows[i].cell[1].text
        question.save()
        questions.append(question)
    print(questions)
    return questions
def get_image_name(text):
    pattern = r"\[file:(.*?)\]"
    matches = re.findall(pattern, text)
    return matches